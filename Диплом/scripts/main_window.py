import sys
from PyQt5 import QtWidgets
from PyQt5.QtCore import*
from PyQt5.QtWidgets import QAction
from PyQt5.QtGui import QFont, QColor, QBrush
import datetime
from time import sleep
import sys
import os
import shutil
import threading
import main_desidn_window
import database
import parsing_window
from docx import Document
from docx.shared import Inches, Pt, Mm, Cm

class MainWindow(QtWidgets.QMainWindow, main_desidn_window.Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        # try:

        self.pars_window = parsing_window.Parsing()

        self.insert_data_for_residents()
        self.insert_data_for_storage()
        self.insert_data_for_work_statement()
        self.insert_data_for_housing_stock()
        self.insert_data_for_common_areas()
        self.insert_data_for_debtors()
        self.insert_data_for_suppliers()
        self.insert_data_for_statement()
        self.insert_data_for_utilites()
        self.bttnAccept.clicked.connect(self.addRowForResidents)
        self.bttnAdd.clicked.connect(self.addResidents)
        self.bttnDel.clicked.connect(self.delRowForResidents)
        self.bttnDellHousingStock.clicked.connect(self.delRowForHousingStock)
        self.bttnDellCommonAreas.clicked.connect(self.delRowForCommonAreas)
        self.bttnDelForDebtors.clicked.connect(self.delRowForDebtors)
        self.bttnDelForSuppliers.clicked.connect(self.delRowForSuppliers)
        self.bttnDelForStatement.clicked.connect(self.delRowForStatement)
        self.bttnAddForStorage.clicked.connect(self.AddForStorage)
        self.bttnDelForStorage.clicked.connect(self.delete_data_for_storage)
        self.bttnRefreshForStorage.clicked.connect(self.insert_data_for_storage)
        self.listWidgetForStorage.itemDoubleClicked.connect(self.openFileForStorage)
        self.listWidgetForCheckable.itemDoubleClicked.connect(self.openFileForStatement)
        self.treeForDocs.itemDoubleClicked.connect(self.treeDocs)
        self.bttnAddForStorageArrival.clicked.connect(self.addForWorkStatement)
        self.bttnAccept_2.clicked.connect(self.addRowForHousingStock)
        self.bttnAccept_3.clicked.connect(self.addRowForCommonAreas)
        self.bttnAccept_7.clicked.connect(self.addRowForDebtors)
        self.bttnAccept_8.clicked.connect(self.addRowForSuppliers)
        self.bttnAddHousingStock.clicked.connect(self.addHousingStock)
        self.bttnAddCommonAreas.clicked.connect(self.addCommonAreas)
        self.bttnAddForDebtors.clicked.connect(self.addDebtors)
        self.bttnAddForSuppliers.clicked.connect(self.addSuppliers)
        self.bttnAcceptFotStatement.clicked.connect(self.addStatement)
        self.statusBar().setFont(QFont("Times New Roman", 14, QFont.Normal))
        self.comboType_2.activated.connect(self.CheckableApplications)
        self.bttnCalculating.clicked.connect(self.Calculating)
        self.bttnParsing.clicked.connect(self.parsing)
        self.bttnSearchForResidents.clicked.connect(self.searchForResidents)
        self.bttnRefreshTableResidents.clicked.connect(self.insert_data_for_residents)
        self.bttnSearchForStorage.clicked.connect(self.searchForStorage)
        self.bttnSearchForHousingStock.clicked.connect(self.searchForHousingStock)
        self.bttnRefreshTableHousingStock.clicked.connect(self.insert_data_for_housing_stock)
        self.bttnSearchForCommonAreas.clicked.connect(self.searchForCommonAreas)
        self.bttnRefreshTableCommonAreas.clicked.connect(self.insert_data_for_common_areas)
        self.bttnRefreshTableDebtors.clicked.connect(self.insert_data_for_debtors)
        self.bttnSearchForDebtors.clicked.connect(self.searchForDebtors)
        self.bttnRefreshTableSuppliers.clicked.connect(self.insert_data_for_suppliers)
        self.bttnSearchForSuppliers.clicked.connect(self.searchForSuppliers)
        self.bttnRefreshTableWorkStatement.clicked.connect(self.insert_data_for_work_statement)
        self.bttnSearchForWorkStatement.clicked.connect(self.searchForWorkStatement)
        self.bttnSearchForStatement.clicked.connect(self.searchForStatement)
        self.bttnRefreshTableStatement.clicked.connect(self.insert_data_for_statement)
        self.tableResidents.cellDoubleClicked.connect(self.selectRowForChangeForResidents)
        self.bttnChangeForResidents.clicked.connect(self.changeForResidents)
        self.tableHousingStock.cellDoubleClicked.connect(self.selectRowForChangeForHousingStock)
        self.bttnChangeHousingStock.clicked.connect(self.changeForHousingStock)
        self.tableCommonAreas.cellDoubleClicked.connect(self.selectRowForChangeForCommonAreas)
        self.bttnChangeCommonAreas.clicked.connect(self.changeForCommonAreas)
        self.tableDebtors.cellDoubleClicked.connect(self.selectRowForChangeForDebtors)
        self.bttnChangeForDebtors.clicked.connect(self.changeForDebtors)
        self.tableSuppliers.cellDoubleClicked.connect(self.selectRowForChangeForSuppliers)
        self.bttnChangeForSuppliers.clicked.connect(self.changeForSuppliers)
        self.bttnSearchForAdvancedSearching.clicked.connect(self.advancedSearching)
        #self.comboForWorkStatement.currentTextChanged.connect(self.outputStatistic)
        self.bttnAcceptForCommonAreaForStatistic.clicked.connect(self.outputStatisticForWorkStatement)
        self.bttnArrival.clicked.connect(self.Arrival)
        self.bttnNotArrival.clicked.connect(self.NotArrival)
        self.bttnWaiting.clicked.connect(self.Waiting)
        self.listForDocs.itemClicked.connect(self.nameDoc1)
        self.bttnForDocOpen.clicked.connect(self.openDoc)
        self.bttnForDocDel.clicked.connect(self.delDoc)
        self.bttnExit.clicked.connect(self.close)
        self.bttnForCreateReference.clicked.connect(self.createReference)
        self.pushButton_2.clicked.connect(self.createStorageP)
        self.pushButton_5.clicked.connect(self.createStorageU)
        self.bttnPrintForCommonArea.clicked.connect(self.creareServices)
        self.bttnAceptForStatistic.clicked.connect(self.aceptForStatistic)
        self.comboForWorkStatement_2.currentIndexChanged.connect(self.hideCombo)

        # except:
        #     print("Ошибка в инициализации функций")

        def close(self):
            self.close


        #try:

        self.tableAdvancedSearching.setColumnWidth(0, 300)
        self.tableAdvancedSearching.setColumnWidth(1, 300)

        #self.directory2 = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/serviceCommonArea")

        database.cur.execute("SELECT Наименование FROM 'common areas';")
        self.result_name_common_areas = database.cur.fetchall()
        for i in range(len(self.result_name_common_areas)):
            self.comboForCommonAreaForPlaceForStatistic.addItems(self.result_name_common_areas[i])

        self.comboForWorkStatement.clear()
        database.cur.execute("SELECT Тип_работ FROM 'works statement';")
        self.result_type_work_statement = database.cur.fetchall()
        print(self.result_type_work_statement)
        for i in range(len(self.result_type_work_statement)):
            self.comboForWorkStatement.addItems(self.result_type_work_statement[i])

        self.directory5 = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/serviceCommonArea")

        self.listWidgetForCommonArea.clear()

        for i in range(len(self.directory5)):
            self.listWidgetForCommonArea.insertItem(0, str(self.directory5[i]))

        database.cur.execute("SELECT ФИО FROM residents;")
        self.result_name_for_searching = database.cur.fetchall()
        for i in range(len(self.result_name_for_searching)):
            self.comboForAdvancedSearching.addItems(self.result_name_for_searching[i])

        database.cur.execute("SELECT [Наименование] FROM 'common areas';")
        self.result_type_common_area = database.cur.fetchall()
        for i in range(len(self.result_type_common_area)):
            self.comboForCommonArea.addItems(self.result_type_common_area[i])

        database.cur.execute("SELECT [Наименование работ] FROM 'type of work';")
        self.result_type = database.cur.fetchall()
        for i in range(len(self.result_type)):
            self.comboType.addItems(self.result_type[i])

        #add combo name from residents
        database.cur.execute("SELECT [ФИО] FROM 'residents';")
        self.result_name = database.cur.fetchall()
        for i in range(len(self.result_name)):
            self.comboName.addItems(self.result_name[i])

        database.cur.execute("SELECT [Наименование] FROM 'storage'")
        self.result_name = database.cur.fetchall()
        for i in range(len(self.result_name)):
            self.comboForNameStorageOutgoing.addItems(self.result_name[i])

        directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/work")

        for i in range(len(directory)):
            self.listWidgetForStorage.insertItem(0, str(directory[i]))

        self.directory1 = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/completed")

        for i in range(len(self.directory1)):
            self.listWidgetForCheckable.insertItem(0, str(self.directory1[i]))

        self.directory2 = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/storage/P")
        #listWidgetForStorageForDocumentP

        for i in range(len(self.directory2)):
            self.listWidgetForStorageForDocumentP.insertItem(0, str(self.directory2[i]))

        self.directory3 = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/storage/U")
        #listWidgetForStorageForDocumentP

        for i in range(len(self.directory3)):
            self.listWidgetForStorageForDocumentU.insertItem(0, str(self.directory3[i]))

        #except:
            #print("Ошибка в наполнении виджетов")


        try:

            timer = QTimer(self)
            timer.timeout.connect(self.showTime)
            timer.start(1000)

        except:
            print("Ошибка в таймере")

    def hideCombo(self, index):
        print(index)
        if index == 0:
            self.dateEdit_3.show()
        elif index == 1:
            self.dateEdit_3.hide()

    def aceptForStatistic(self):
        print(self.comboForWorkStatement.currentText())
        print(self.comboForWorkStatement_2.currentText())
        date = '{}.{}.{}'.format(self.dateEdit_3.dateTime().toString('dd'), self.dateEdit_3.dateTime().toString('MM'), self.dateEdit_3.dateTime().toString('yyyy'))

        print(date)


        if self.comboForWorkStatement_2.currentText() == "Выполнено":

            database.cur.execute("UPDATE 'works statement' SET Статус = ?, Дата_окончания = ? WHERE Тип_работ = ?;", (self.comboForWorkStatement_2.currentText(), date, self.comboForWorkStatement.currentText(),))
            database.conn.commit()

        elif self.comboForWorkStatement_2.currentText() == "Не выполнено":
            date = "-"
            database.cur.execute("UPDATE 'works statement' SET Статус = ?, Дата_окончания = ? WHERE Тип_работ = ?;", (self.comboForWorkStatement_2.currentText(), date, self.comboForWorkStatement.currentText(),))
            database.conn.commit()


    def creareServices(self):

        #try:

            print(self.comboForCommonArea.currentText(), " - Помещение")
            print(self.lineEdit_5.text(), " - Электричество")
            print(self.lineEdit_6.text(), " - Вода")
            print(self.lineEdit_7.text(), " - Водоотведение")
            print(self.SumCommonArea.text(), " - Итого")

            desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
            document = Document()

            style = document.styles["Normal"]
            style.font.size = Pt(14)
            style.font.name = "Times New Roman"

            #ЗАГОЛОВОК
            paragraphForHead = document.add_paragraph()
            runForHead = paragraphForHead.add_run("СУиРОИ")
            runForHead.bold = True
            runForHead.font.size = Pt(22)
            paragraphForHead.alignment = 1
            paragraphForBottom1 = document.add_paragraph()
            runForBottom = paragraphForBottom1.add_run(str(datetime.datetime.now().strftime('%m.%d.%Y')) + " г.")
            paragraphForBottom1.alignment = 2
            paragraphForBottom1.bold = True

            #ТЕКСТ
            paragraphForText1 = document.add_paragraph()
            runForText1 = paragraphForText1.add_run("Помещение - " + str(self.comboForCommonArea.currentText()))
            paragraphForText2 = document.add_paragraph()
            runForText2 = paragraphForText2.add_run("Электричество - " + str(self.lineEdit_5.text()) + " кВт.")
            paragraphForText3 = document.add_paragraph()
            runForText3 = paragraphForText3.add_run("Вода - " + str(self.lineEdit_6.text()) + " м.куб.")
            paragraphForText4 = document.add_paragraph()
            runForText4 = paragraphForText4.add_run("Водоотведение - " + str(self.lineEdit_7.text()) + " м.куб.")
            paragraphForText5 = document.add_paragraph()
            runForText5 = paragraphForText5.add_run("Итого - " + str(self.SumCommonArea.text()))

            #НИЖНЯЯ ЧАСТЬ
            paragraphForBottom = document.add_paragraph()
            runForBottom = paragraphForBottom.add_run("Директор ООО «УК ЭЖК КОЛОСКОВА»                          Ахмедова Г.Г.")



            document.save(desktop + "/" + "диплом vol.2/docs/serviceCommonArea" + "/" + str(self.comboForCommonArea.currentText()) + " " + str(datetime.datetime.now().strftime('%Y.%m.%d')) + ".docx")

            print("save")

            self.listWidgetForCommonArea.clear()

            self.directory5 = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/serviceCommonArea")

            for i in range(len(self.directory5)):
                self.listWidgetForCommonArea.insertItem(0, str(self.directory5[i]))
        
        #except:
            #print("Ошибка в создании документа СУиРОИ")

    def createStorageU(self):
        try:
            print(self.comboForNameStorageOutgoing.currentText(), " - Наименование")
            print(self.lineForCountStorageOutgoing.text(), " - Количество")

            desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
            document = Document()

            style = document.styles["Normal"]
            style.font.size = Pt(14)
            style.font.name = "Times New Roman"

            #ЗАГОЛОВОК
            paragraphForHead = document.add_paragraph()
            runForHead = paragraphForHead.add_run("Уход инвентаря")
            runForHead.bold = True
            runForHead.font.size = Pt(22)
            paragraphForHead.alignment = 1

            #ТЕКСТ
            paragraphForText = document.add_paragraph()
            runForText = paragraphForText.add_run("Инвентарь " + str(self.comboForNameStorageOutgoing.currentText()) + " в количестве " + str(self.lineForCountStorageOutgoing.text()))

            #НИЖНЯЯ ЧАСТЬ
            paragraphForBottom = document.add_paragraph()
            runForBottom = paragraphForBottom.add_run(str(datetime.datetime.now().strftime('%Y.%m.%d')))
            runForBottom = paragraphForBottom.add_run("                                                     Подпись __________________")



            document.save(desktop + "/" + "диплом vol.2/docs/storage/U" + "/" + str(self.comboForNameStorageOutgoing.currentText()) + " " + str(datetime.datetime.now().strftime('%Y.%m.%d')) + ".docx")

            self.listWidgetForStorageForDocumentU.clear()

            self.directory3 = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/storage/U")
            #listWidgetForStorageForDocumentU
        
            for i in range(len(self.directory3)):
                self.listWidgetForStorageForDocumentU.insertItem(0, str(self.directory3[i]))

        except:
            print("Ошибка в создании документа ухода товара")

    def createStorageP(self):
        #try:
            print(self.lineForNameStorageArrival.text(), " - Наименование")
            print(self.lineForCountStorageArrival.text(), " - Количество")
            print(self.lineForPriceStorageArrival.text(), " - Цена")

            desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
            document = Document()

            style = document.styles["Normal"]
            style.font.size = Pt(14)
            style.font.name = "Times New Roman"

            #ЗАГОЛОВОК
            paragraphForHead = document.add_paragraph()
            runForHead = paragraphForHead.add_run("Поступление инвентаря от " + str(datetime.datetime.now().strftime('%Y.%m.%d')))
            runForHead.bold = True
            runForHead.font.size = Pt(22)
            paragraphForHead.alignment = 1

            #ТЕКСТ
            paragraphForHead = document.add_paragraph()
            runForHead = paragraphForHead.add_run("")
            table = document.add_table(rows = 1, cols = 5, style = "Table Grid")
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = 'Товар'
            hdr_cells[2].text = 'Количество'
            hdr_cells[3].text = 'Цена'
            hdr_cells[4].text = 'Сумма'

            row_cells = table.add_row().cells

            row_cells[0].text = str(1)
            row_cells[1].text = str(self.lineForNameStorageArrival.text())
            row_cells[2].text = str(self.lineForCountStorageArrival.text())
            row_cells[3].text = str(self.lineForPriceStorageArrival.text())
            row_cells[4].text = str(int(self.lineForPriceStorageArrival.text()) * int(self.lineForPriceStorageArrival.text()))

            paragraphForHead = document.add_paragraph()
            runForHead = paragraphForHead.add_run("")

            #НИЖНЯЯ ЧАСТЬ
            paragraphForBottom = document.add_paragraph()
            runForBottom = paragraphForBottom.add_run(str(datetime.datetime.now().strftime('%Y.%m.%d')))
            runForBottom = paragraphForBottom.add_run("                                                     Подпись __________________")



            document.save(desktop + "/" + "диплом vol.2/docs/storage/P" + "/" + str(self.lineForNameStorageArrival.text()) + " " + str(datetime.datetime.now().strftime('%Y.%m.%d')) + ".docx")

            self.listWidgetForStorageForDocumentP.clear()

            self.directory2 = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/storage/P")
            #listWidgetForStorageForDocumentP

            
            for i in range(len(self.directory2)):
                self.listWidgetForStorageForDocumentP.insertItem(0, str(self.directory2[i]))

        #except:
            #print("Ошибка в создании документа прихода товара")

        # self.listWidgetForCheckable.clear()
        # directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/waiting")
        # for i in range(len(directory)):
        #     self.listWidgetForCheckable.insertItem(0, str(directory[i]))

    def createReference(self):
        try:
            print(self.lineNameForAdvancedSearching.text(), " - ФИО")
            print(self.lineGenderForAdvancedSearching.text(), " - Пол")
            print(self.lineArmyForAdvancedSearching.text(), " - Воинская обязанность")
            print(self.linePasportForAdvancedSearching.text(), " - Паспортные данные")
            print(self.lineSnilsForAdvancedSearching.text(), " - СНИЛС")
            print(self.linePodezdForAdvancedSearching.text(), " - № подъезда")
            print(self.lineKvartiraForAdvancedSearching.text(), " - № квартиры")
            print(self.linePlaceForAdvancedSearching.text(), " - Площадь помещения м2")
            print(self.lineLivePlaceForAdvancedSearching.text(), " - Площадь жилого помещения м2")
            print(self.lineRegistrationForAdvancedSearching.text(), " - Регистрация")
            print(self.lineDebtorsForAdvancedSearching.text(), " - Задолженность")
            print(self.linePriceForAdvancedSearching.text(), " - Цена руб")

            desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')

            document = Document()

            style = document.styles["Normal"]
            style.font.size = Pt(14)
            style.font.name = "Times New Roman"

            #section = document.sections[-1] 
            #section.left_margin = Inches(1.5)

            #para.paragraph_format.left_indent = Inches( 0.5 )

            paragraphForHead = document.add_paragraph()
            runForHead = paragraphForHead.add_run("Справка")
            runForHead.bold = True
            runForHead.font.size = Pt(22)
            paragraphForHead.alignment = 1

            paragraphForBottom1 = document.add_paragraph()
            runForBottom1 = paragraphForBottom1.add_run(str(datetime.datetime.now().strftime('%m.%d.%Y')) + " г.")
            paragraphForBottom1.alignment = 2
            paragraphForBottom1.bold = True

            paragraphForText1 = document.add_paragraph()
            runForText = paragraphForText1.add_run("Справка дана гр. " + str(self.lineNameForAdvancedSearching.text()))
            runForText = paragraphForText1.add_run(" в том, что она действительно проживает по адресу г. Калининград, ул. Сержанта Колоскова 8" 
                                                    + ", зарегестрирована " + str(self.lineRegistrationForAdvancedSearching.text()) + ".")
            paragraphForText2 = document.add_paragraph()
            runForText1 = paragraphForText2.add_run("Жилая площадь квартиры " + str(self.lineLivePlaceForAdvancedSearching.text()) + " кв.м." 
                                                    + ", подъезд " + str(self.linePodezdForAdvancedSearching.text()) + ", квартира " 
                                                    + str(self.lineKvartiraForAdvancedSearching.text()) + ".")

            paragraphForText3 = document.add_paragraph("Задолженности " + str(self.lineDebtorsForAdvancedSearching.text()) 
                                                    + ", итого " + str(self.linePriceForAdvancedSearching.text()) + ".")
            runForText2 = paragraphForText3.add_run()

            paragraphForBottom = document.add_paragraph()
            runForBottom = paragraphForBottom.add_run("Директор ООО «УК ЭЖК КОЛОСКОВА»                         Ахмедова Г.Г.")

            document.save(desktop + "/" + "диплом vol.2/docs/forResidents" + "/" + str(self.lineNameForAdvancedSearching.text()) + ".docx")

        except:
            print("Ошибка в создании документа справка")

    def createStatementDocument(self, resident ,name, data):
        try:
            desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
            print(resident)
            print(name)
            print(data)
            document = Document()

            style = document.styles["Normal"]
            style.font.size = Pt(14)
            style.font.name = "Times New Roman"

            #ЗАГОЛОВОК
            paragraphForHead = document.add_paragraph()
            runForHead = paragraphForHead.add_run("Заявка на проведение работ")
            runForHead.bold = True
            runForHead.font.size = Pt(22)
            paragraphForHead.alignment = 1
            paragraphForBottom1 = document.add_paragraph()
            runForBottom1 = paragraphForBottom1.add_run(str(datetime.datetime.now().strftime('%m.%d.%Y')) + " г.")
            paragraphForBottom1.alignment = 2
            paragraphForBottom1.bold = True

            #ТЕКСТ
            paragraphForText = document.add_paragraph()
            runForText = paragraphForText.add_run(str(name) + " будет проводиться " + str(data) + ".")

            #НИЖНЯЯ ЧАСТЬ
            paragraphForBottom = document.add_paragraph()
            runForBottom = paragraphForBottom.add_run("Директор ООО «УК ЭЖК КОЛОСКОВА»                          Ахмедова Г.Г.")



            document.save(desktop + "/" + "диплом vol.2/docs/request/waiting" + "/" + str(resident) + " " + str(data) + ".docx")

            self.listWidgetForCheckable.clear()
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/waiting")
            for i in range(len(directory)):
                self.listWidgetForCheckable.insertItem(0, str(directory[i]))

        except:
            print("Ошибка в создании документа заявки")

    def createWorkStatementDocument(self, name, data):
        try:
            desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
            print(name)
            print(data)
            document = Document()

            style = document.styles["Normal"]
            style.font.size = Pt(14)
            style.font.name = "Times New Roman"

            #ЗАГОЛОВОК
            paragraphForHead = document.add_paragraph()
            runForHead = paragraphForHead.add_run("Заявка на проведение ремонтных работ")
            runForHead.bold = True
            runForHead.font.size = Pt(22)
            paragraphForHead.alignment = 1
            paragraphForBottom1 = document.add_paragraph()
            runForBottom1 = paragraphForBottom1.add_run(str(datetime.datetime.now().strftime('%m.%d.%Y')) + " г.")
            paragraphForBottom1.alignment = 2
            paragraphForBottom1.bold = True

            #ТЕКСТ
            paragraphForText1 = document.add_paragraph()
            runForText1 = paragraphForText1.add_run(str(name) + " будет проводиться " + str(data) + ".")

            #НИЖНЯЯ ЧАСТЬ
            paragraphForBottom = document.add_paragraph()
            runForBottom = paragraphForBottom.add_run("Директор ООО «УК ЭЖК КОЛОСКОВА»                          Ахмедова Г.Г.")

            document.save(desktop + "/" + "диплом vol.2/docs/work" + "/" + str(name) + ".docx")
            self.listWidgetForStorage.clear()
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/work")
            for i in range(len(directory)):
                self.listWidgetForStorage.insertItem(0, str(directory[i]))

        except:
            print("Ошибка в создании документа заявки на проведение работ")

    def delDoc(self):
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        try:
            if self.treeForDocs.currentItem().text(0) == "Выполненные":
                os.remove(desktop + "/" + "диплом vol.2/docs/request/completed" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Невыполненные":
                os.remove(desktop + "/" + "диплом vol.2/docs/request/unfulfilled" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Ожидающие":
                os.remove(desktop + "/" + "диплом vol.2/docs/request/waiting" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Проведение работ":
                os.remove(desktop + "/" + "диплом vol.2/docs/work" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "СУиРОИ":
                os.remove(desktop + "/" + "диплом vol.2/docs/serviceCommonArea" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Жильцам":
                os.remove(desktop + "/" + "диплом vol.2/docs/forResidents/natural person" + "/" + nameDoc)

        except:
            pass

    def openDoc(self):  
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        try:
            if self.treeForDocs.currentItem().text(0) == "Выполненные":
                os.startfile(desktop + "/" + "диплом vol.2/docs/request/completed" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Невыполненные":
                os.startfile(desktop + "/" + "диплом vol.2/docs/request/unfulfilled" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Ожидающие":
                os.startfile(desktop + "/" + "диплом vol.2/docs/request/waiting" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Проведение работ":
                os.startfile(desktop + "/" + "диплом vol.2/docs/work" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "СУиРОИ":
                os.startfile(desktop + "/" + "диплом vol.2/docs/serviceCommonArea" + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Жильцам":
                os.startfile(desktop + "/" + "диплом vol.2/docs/forResidents" + "/" + nameDoc)

        except:
            pass

    def Waiting(self):
        try:
            if self.treeForDocs.currentItem().text(0) == "Выполненные":
                old = "диплом vol.2/docs/request/completed"
                new = "диплом vol.2/docs/request/waiting"
                print(nameDoc)
                print("Невыполненные")
                desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
                print(desktop + "/" + old + "/" + nameDoc)
                shutil.move(desktop + "/" + old + "/" + nameDoc, desktop + "/" + new + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Невыполненные":
                old = "диплом vol.2/docs/request/unfulfilled"
                new = "диплом vol.2/docs/request/waiting"
                print(nameDoc)
                print("Невыполненные")
                desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
                print(desktop + "/" + old + "/" + nameDoc)
                shutil.move(desktop + "/" + old + "/" + nameDoc, desktop + "/" + new + "/" + nameDoc)
        except:
            pass

    def NotArrival(self):
        try:
            if self.treeForDocs.currentItem().text(0) == "Ожидающие":
                old = "диплом vol.2/docs/request/waiting"
                new = "диплом vol.2/docs/request/unfulfilled"
                print(nameDoc)
                print("Невыполненные")
                desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
                print(desktop + "/" + old + "/" + nameDoc)
                shutil.move(desktop + "/" + old + "/" + nameDoc, desktop + "/" + new + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Выполненные":
                old = "диплом vol.2/docs/request/completed"
                new = "диплом vol.2/docs/request/unfulfilled"
                print(nameDoc)
                print("Невыполненные")
                desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
                print(desktop + "/" + old + "/" + nameDoc)
                shutil.move(desktop + "/" + old + "/" + nameDoc, desktop + "/" + new + "/" + nameDoc)
        except:
            pass
        
    def Arrival(self):
        try:
            if self.treeForDocs.currentItem().text(0) == "Ожидающие":
                old = "диплом vol.2/docs/request/waiting"
                new = "диплом vol.2/docs/request/completed"
                print(nameDoc)
                print("Выполненные")
                desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
                print(desktop + "/" + old + "/" + nameDoc)
                shutil.move(desktop + "/" + old + "/" + nameDoc, desktop + "/" + new + "/" + nameDoc)

            elif self.treeForDocs.currentItem().text(0) == "Невыполненные":
                old = "диплом vol.2/docs/request/unfulfilled"
                new = "диплом vol.2/docs/request/completed"
                print(nameDoc)
                print("Невыполненные")
                desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
                print(desktop + "/" + old + "/" + nameDoc)
                shutil.move(desktop + "/" + old + "/" + nameDoc, desktop + "/" + new + "/" + nameDoc)
        except:
            pass

    def nameDoc1(self, item):
        global nameDoc
        nameDoc = item.text()

    def outputStatisticForWorkStatement(self, index):

        try:
            database.cur.execute("SELECT Электричество, Вода, Водоотведение, Итого FROM utilities WHERE Помещение = ? AND Дата = ?;", (self.comboForCommonAreaForPlaceForStatistic.currentText(), self.lineEditDataForStatistic.text(),))
            dataForStatistc = database.cur.fetchall()

            database.cur.execute("SELECT Площадь_м2 FROM 'common areas' WHERE Наименование = ?;", (self.comboForCommonAreaForPlaceForStatistic.currentText(),))
            SquareForStatistc = database.cur.fetchall()

            self.lineEditForElectricForStatistic.setText(str(dataForStatistc[0][0]))
            self.lineEditForWaterForStatistic.setText(str(dataForStatistc[0][1]))
            self.lineEditForWater2ForStatistic.setText(str(dataForStatistc[0][2]))
            self.lineEditForSquareForStatistic.setText(str(SquareForStatistc[0][0]))
            self.labelForSum.setText(str(dataForStatistc[0][3]) + " руб.")

        except:
            print("Нету")

    # def outputStatistic(self, index):
    #     database.cur.execute("SELECT Потрачечно, Дата_окончания FROM 'works statement' WHERE  Тип_работ = ?;", (index,))
    #     work_statement_option = database.cur.fetchall()
    #     self.lineEditWastedForWorkStatement.setText(str(work_statement_option[0][0]) + " Инструментов")
    #     self.lineEditDataEndForWorkStatement.setText(str(work_statement_option[0][1]))

    def advancedSearching(self):
        residentsData = []
        housingStockData = []
        debtorsData = []
        statementData = []
        database.cur.execute("SELECT ФИО, Пол, Воинская_обязанность, Паспортные_данные, СНИЛС, №_подъезда, №_квартиры FROM residents WHERE ФИО = ?;", (self.comboForAdvancedSearching.currentText(),))
        residentsData.append(database.cur.fetchone())
        self.lineNameForAdvancedSearching.setText(str(residentsData[0][0]))
        self.lineGenderForAdvancedSearching.setText(str(residentsData[0][1]))
        self.lineArmyForAdvancedSearching.setText(str(residentsData[0][2]))
        self.linePasportForAdvancedSearching.setText(str(residentsData[0][3]))
        self.lineSnilsForAdvancedSearching.setText(str(residentsData[0][4]))
        self.linePodezdForAdvancedSearching.setText(str(residentsData[0][5]))
        self.lineKvartiraForAdvancedSearching.setText(str(residentsData[0][6]))
        database.cur.execute("SELECT Площадь_помещения_м2, Площадь_жилого_помещения_м2, Регистрация FROM 'housing stock' WHERE Владелец = ?;", (self.comboForAdvancedSearching.currentText(),))
        housingStockData.append(database.cur.fetchone())
        self.linePlaceForAdvancedSearching.setText(str(housingStockData[0][0]))
        self.lineLivePlaceForAdvancedSearching.setText(str(housingStockData[0][1]))
        self.lineRegistrationForAdvancedSearching.setText(str(housingStockData[0][2]))
        database.cur.execute("SELECT Тип_задолженности, Цена_руб FROM debtors WHERE ФИО_должника = ?;", (self.comboForAdvancedSearching.currentText(),))
        debtorsData.append(database.cur.fetchone())
        try:
            self.lineDebtorsForAdvancedSearching.setText(str(debtorsData[0][0]))
            self.linePriceForAdvancedSearching.setText(str(debtorsData[0][1]))
        except:
            self.lineDebtorsForAdvancedSearching.setText(str("-"))
            self.linePriceForAdvancedSearching.setText(str("-"))

        database.cur.execute("SELECT Тип_работ, Дата_проведения FROM statement WHERE ФИО_заявителя = ?;", (self.comboForAdvancedSearching.currentText(),))
        statementData = database.cur.fetchall()

        #print(statementData)
        #self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))

        if statementData == []:
            self.tableAdvancedSearching.clearContents()
            self.tableAdvancedSearching.setRowCount(0)
        else:
            self.tableAdvancedSearching.verticalHeader().setVisible(False)
            for i in range(len(statementData)):

                self.tableAdvancedSearching.setRowCount(len(statementData))

                for i in range(len(statementData)):
                    self.tableAdvancedSearching.setItem(i, 0, QtWidgets.QTableWidgetItem(str(statementData[i][0])))
                    if i == 0:
                        self.tableAdvancedSearching.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                    elif i % 2 == 0:
                        self.tableAdvancedSearching.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

                for i in range(len(statementData)):
                    self.tableAdvancedSearching.setItem(i, 1, QtWidgets.QTableWidgetItem(str(statementData[i][1])))
                    if i == 0:
                        self.tableAdvancedSearching.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                    elif i % 2 == 0:
                        self.tableAdvancedSearching.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

    def selectRowForChangeForSuppliers(self):
        global oldSuppliers
        row = self.tableSuppliers.currentRow()
        oldSuppliers = []
        for i in range(self.tableSuppliers.columnCount()):
            oldSuppliers.append(self.tableSuppliers.item(row, i).text())

    def changeForSuppliers(self):
        row = self.tableSuppliers.currentRow()
        rowForDb = self.tableSuppliers.currentRow() + 1

        newSuppliers = []
        for i in range(self.tableSuppliers.columnCount()):
            newSuppliers.append(self.tableSuppliers.item(row, i).text())
        newSuppliers.append(rowForDb)

        database.cur.executemany("""UPDATE 'suppliers' SET ID = ?, Наименование = ?, Адрес = ?, ИНН = ?, КПП = ?, Р_счёт = ?, Корр_счёт = ?, БИК = ?, Телефон = ?, ФИО = ? WHERE ID = ?""", (newSuppliers,))

        database.conn.commit()

    def selectRowForChangeForDebtors(self):
        global oldDebtors
        row = self.tableDebtors.currentRow()
        oldDebtors = []
        for i in range(self.tableDebtors.columnCount()):
            oldDebtors.append(self.tableDebtors.item(row, i).text())

    def changeForDebtors(self):
        row = self.tableDebtors.currentRow()
        rowForDb = self.tableDebtors.currentRow() + 1

        newDebtors = []
        for i in range(self.tableDebtors.columnCount()):
            newDebtors.append(self.tableDebtors.item(row, i).text())
        newDebtors.append(rowForDb)

        database.cur.executemany("""UPDATE 'debtors' SET ID = ?, ФИО_должника = ?, Тип_задолженности = ?, Цена_руб = ? WHERE ID = ?""", (newDebtors,))

        database.conn.commit()

    def selectRowForChangeForCommonAreas(self):
        global oldCommonAreas
        row = self.tableCommonAreas.currentRow()
        oldCommonAreas = []
        for i in range(self.tableCommonAreas.columnCount()):
            oldCommonAreas.append(self.tableCommonAreas.item(row, i).text())

    def changeForCommonAreas(self):
        row = self.tableCommonAreas.currentRow()
        rowForDb = self.tableCommonAreas.currentRow() + 1

        newCommonAreas = []
        for i in range(self.tableCommonAreas.columnCount()):
            newCommonAreas.append(self.tableCommonAreas.item(row, i).text())
        newCommonAreas.append(rowForDb)

        database.cur.executemany("""UPDATE 'common areas' SET ID = ?, Наименование = ?, Площадь_м2 = ? WHERE ID = ?""", (newCommonAreas,))

        database.conn.commit()

    def selectRowForChangeForHousingStock(self):
        global oldHousingStock
        row = self.tableHousingStock.currentRow()
        oldHousingStock = []
        for i in range(self.tableHousingStock.columnCount()):
            oldHousingStock.append(self.tableHousingStock.item(row, i).text())

    def changeForHousingStock(self):
        row = self.tableHousingStock.currentRow()
        rowForDb = self.tableHousingStock.currentRow() + 1

        newHousingStock = []
        for i in range(self.tableHousingStock.columnCount()):
            newHousingStock.append(self.tableHousingStock.item(row, i).text())
        newHousingStock.append(rowForDb)

        database.cur.executemany("""UPDATE 'housing stock' SET ID = ?, Владелец = ?, Площадь_помещения_м2 = ?, Площадь_жилого_помещения_м2 = ?, Регистрация = ? WHERE ID = ?""", (newHousingStock,))

        database.conn.commit()


    def selectRowForChangeForResidents(self):
        global oldResident
        try:
            row = self.tableResidents.currentRow()
            oldResident = []
            for i in range(self.tableResidents.columnCount()):
                oldResident.append(self.tableResidents.item(row, i).text())
        except:
            print("Ошибка в дабл клике по таблице жильцы")

    def changeForResidents(self):#протестить до конца изменение строки, выглядит как нездоровая хуйня
        row = self.tableResidents.currentRow()
        rowForDb = self.tableResidents.currentRow() + 1

        newResident = []
        for i in range(self.tableResidents.columnCount()):
            newResident.append(self.tableResidents.item(row, i).text())
        newResident.append(rowForDb)

        database.cur.executemany("""UPDATE residents SET ID = ?, ФИО = ?, Пол = ?, Воинская_обязанность = ?, Паспортные_данные = ?, СНИЛС = ?, №_подъезда = ?, №_квартиры = ?, Задолженность = ? WHERE ID = ?""", (newResident,))

        database.conn.commit()

    def searchForStatement(self):
        valueForStatement = self.lineForSearchingForStatement.text()
        print(valueForStatement)

        try:
            newValueForTableStatement = []
            database.cur.execute("SELECT * FROM 'statement' WHERE ФИО_заявителя = ?", ([valueForStatement]))
            newValueForTableStatement = database.cur.fetchall()
            print(newValueForTableStatement)
            if newValueForTableStatement == []:
                print("Этот не ФИО заявителя")
            else:
                self.tableStatement.clearContents()
                self.tableStatement.setRowCount(len(newValueForTableStatement))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][0])))
                if i == 0:
                    self.tableStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][1])))
                if i == 0:
                    self.tableStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][2])))
                if i == 0:
                    self.tableStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][3])))
                if i == 0:
                    self.tableStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableStatement = []
            database.cur.execute("SELECT * FROM 'statement' WHERE Тип_работ = ?", ([valueForStatement]))
            newValueForTableStatement = database.cur.fetchall()
            print(newValueForTableStatement)
            if newValueForTableStatement == []:
                print("Этот не Тип работ")
            else:
                self.tableStatement.clearContents()
                self.tableStatement.setRowCount(len(newValueForTableStatement))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][0])))
                if i == 0:
                    self.tableStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][1])))
                if i == 0:
                    self.tableStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][2])))
                if i == 0:
                    self.tableStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][3])))
                if i == 0:
                    self.tableStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableStatement = []
            database.cur.execute("SELECT * FROM 'statement' WHERE Дата_проведения = ?", ([valueForStatement]))
            newValueForTableStatement = database.cur.fetchall()
            print(newValueForTableStatement)
            if newValueForTableStatement == []:
                print("Этот не Дата проведения")
            else:
                self.tableStatement.clearContents()
                self.tableStatement.setRowCount(len(newValueForTableStatement))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][0])))
                if i == 0:
                    self.tableStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][1])))
                if i == 0:
                    self.tableStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][2])))
                if i == 0:
                    self.tableStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStatement)):
                self.tableStatement.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableStatement[i][3])))
                if i == 0:
                    self.tableStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

    def searchForWorkStatement(self):
        valueForWorksStatement = self.lineForSearchingForWorkStatement.text()
        print(valueForWorksStatement)

        try:
            newValueForTableWorksStatement = []
            database.cur.execute("SELECT * FROM 'works statement' WHERE Тип_работ = ?", ([valueForWorksStatement]))
            newValueForTableWorksStatement = database.cur.fetchall()
            print(newValueForTableWorksStatement)
            if newValueForTableWorksStatement == []:
                print("Этот не Тип работ")
            else:
                self.tableWorkStatement.clearContents()
                self.tableWorkStatement.setRowCount(len(newValueForTableWorksStatement))

            for i in range(len(newValueForTableWorksStatement)):
                self.tableWorkStatement.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableWorksStatement[i][0])))
                if i == 0:
                    self.tableWorkStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableWorksStatement)):
                self.tableWorkStatement.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableWorksStatement[i][1])))
                if i == 0:
                    self.tableWorkStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableWorksStatement)):
                self.tableWorkStatement.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableWorksStatement[i][2])))
                if i == 0:
                    self.tableWorkStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableWorksStatement = []
            database.cur.execute("SELECT * FROM 'works statement' WHERE Дата_проведения = ?", ([valueForWorksStatement]))
            newValueForTableWorksStatement = database.cur.fetchall()
            print(newValueForTableWorksStatement)
            if newValueForTableWorksStatement == []:
                print("Этот не Дата проведения")
            else:
                self.tableWorkStatement.clearContents()
                self.tableWorkStatement.setRowCount(len(newValueForTableWorksStatement))

            for i in range(len(newValueForTableWorksStatement)):
                self.tableWorkStatement.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableWorksStatement[i][0])))
                if i == 0:
                    self.tableWorkStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableWorksStatement)):
                self.tableWorkStatement.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableWorksStatement[i][1])))
                if i == 0:
                    self.tableWorkStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableWorksStatement)):
                self.tableWorkStatement.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableWorksStatement[i][2])))
                if i == 0:
                    self.tableWorkStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

    def searchForSuppliers(self):
        valueForSuppliers = self.lineForSearchingForSuppliers.text()
        print(valueForSuppliers)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE Наименование = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не Наименование")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE Адрес = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не Адрес")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE ИНН = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не ИНН")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE КПП = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не КПП")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE Р_счёт = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не Р_счёт")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE Корр_счёт = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не Корр_счёт")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE БИК = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не БИК")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE Телефон = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не Телефон")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableSuppliers = []
            database.cur.execute("SELECT * FROM 'suppliers' WHERE ФИО = ?", ([valueForSuppliers]))
            newValueForTableSuppliers = database.cur.fetchall()
            print(newValueForTableSuppliers)
            if newValueForTableSuppliers == []:
                print("Этот не ФИО")
            else:
                self.tableSuppliers.clearContents()
                self.tableSuppliers.setRowCount(len(newValueForTableSuppliers))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableSuppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(newValueForTableSuppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

    def searchForDebtors(self):
        valueForDebtors = self.lineForSearchingForDebtors.text()
        print(valueForDebtors)

        try:
            newValueForTableDebtors = []
            database.cur.execute("SELECT * FROM 'debtors' WHERE ФИО_должника = ?", ([valueForDebtors]))
            newValueForTableDebtors = database.cur.fetchall()
            print(newValueForTableDebtors)
            if newValueForTableDebtors == []:
                print("Этот не ФИО должника")
            else:
                self.tableDebtors.clearContents()
                self.tableDebtors.setRowCount(len(newValueForTableDebtors))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][0])))
                if i == 0:
                    self.tableDebtors.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][1])))
                if i == 0:
                    self.tableDebtors.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][2])))
                if i == 0:
                    self.tableDebtors.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][3])))
                if i == 0:
                    self.tableDebtors.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableDebtors = []
            database.cur.execute("SELECT * FROM 'debtors' WHERE Тип_задолженности = ?", ([valueForDebtors]))
            newValueForTableDebtors = database.cur.fetchall()
            print(newValueForTableDebtors)
            if newValueForTableDebtors == []:
                print("Этот не Тип задолженности")
            else:
                self.tableDebtors.clearContents()
                self.tableDebtors.setRowCount(len(newValueForTableDebtors))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][0])))
                if i == 0:
                    self.tableDebtors.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][1])))
                if i == 0:
                    self.tableDebtors.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][2])))
                if i == 0:
                    self.tableDebtors.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][3])))
                if i == 0:
                    self.tableDebtors.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableDebtors = []
            database.cur.execute("SELECT * FROM 'debtors' WHERE Цена_руб = ?", ([valueForDebtors]))
            newValueForTableDebtors = database.cur.fetchall()
            print(newValueForTableDebtors)
            if newValueForTableDebtors == []:
                print("Этот не Цена")
            else:
                self.tableDebtors.clearContents()
                self.tableDebtors.setRowCount(len(newValueForTableDebtors))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][0])))
                if i == 0:
                    self.tableDebtors.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][1])))
                if i == 0:
                    self.tableDebtors.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][2])))
                if i == 0:
                    self.tableDebtors.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableDebtors)):
                self.tableDebtors.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableDebtors[i][3])))
                if i == 0:
                    self.tableDebtors.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)    

    def searchForCommonAreas(self):
        valueForCommonAreas = self.lineForSearchingForCommonAreas.text()
        print(valueForCommonAreas)

        try:
            newValueForTableCommonAreas = []
            database.cur.execute("SELECT * FROM 'common areas' WHERE Наименование = ?", ([valueForCommonAreas]))
            newValueForTableCommonAreas = database.cur.fetchall()
            print(newValueForTableCommonAreas)
            if newValueForTableCommonAreas == []:
                print("Этот не Наименование")
            else:
                self.tableCommonAreas.clearContents()
                self.tableCommonAreas.setRowCount(len(newValueForTableCommonAreas))

            for i in range(len(newValueForTableCommonAreas)):
                self.tableCommonAreas.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableCommonAreas[i][0])))
                if i == 0:
                    self.tableCommonAreas.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableCommonAreas)):
                self.tableCommonAreas.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableCommonAreas[i][1])))
                if i == 0:
                    self.tableCommonAreas.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableCommonAreas)):
                self.tableCommonAreas.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableCommonAreas[i][2])))
                if i == 0:
                    self.tableCommonAreas.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableCommonAreas = []
            database.cur.execute("SELECT * FROM 'common areas' WHERE Площадь_м2 = ?", ([valueForCommonAreas]))
            newValueForTableCommonAreas = database.cur.fetchall()
            print(newValueForTableCommonAreas)
            if newValueForTableCommonAreas == []:
                print("Этот не Площадь")
            else:
                self.tableCommonAreas.clearContents()
                self.tableCommonAreas.setRowCount(len(newValueForTableCommonAreas))

            for i in range(len(newValueForTableCommonAreas)):
                self.tableCommonAreas.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableCommonAreas[i][0])))
                if i == 0:
                    self.tableCommonAreas.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableCommonAreas)):
                self.tableCommonAreas.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableCommonAreas[i][1])))
                if i == 0:
                    self.tableCommonAreas.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableCommonAreas)):
                self.tableCommonAreas.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableCommonAreas[i][2])))
                if i == 0:
                    self.tableCommonAreas.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

    def searchForHousingStock(self):
        valueForHousingStock = self.lineForSearchingForHousingStock.text()
        print(valueForHousingStock)

        try:
            newValueForTableHousingStock = []
            database.cur.execute("SELECT * FROM 'housing stock' WHERE Владелец = ?", ([valueForHousingStock]))
            newValueForTableHousingStock = database.cur.fetchall()
            print(newValueForTableHousingStock)
            if newValueForTableHousingStock == []:
                print("Этот не Владелец")
            else:
                self.tableHousingStock.clearContents()
                self.tableHousingStock.setRowCount(len(newValueForTableHousingStock))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][0])))
                if i == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][1])))
                if i == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][2])))
                if i == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][3])))
                if i == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][4])))
                if i == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableHousingStock = []
            database.cur.execute("SELECT * FROM 'housing stock' WHERE Площадь_помещения_м2 = ?", ([valueForHousingStock]))
            newValueForTableHousingStock = database.cur.fetchall()
            #print(newValueForTableHousingStock)
            if newValueForTableHousingStock == []:
                print("Этот не Площадь помещения")
            else:
                self.tableHousingStock.clearContents()
                self.tableHousingStock.setRowCount(len(newValueForTableHousingStock))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][0])))
                if i == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][1])))
                if i == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][2])))
                if i == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][3])))
                if i == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][4])))
                if i == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableHousingStock = []
            database.cur.execute("SELECT * FROM 'housing stock' WHERE Площадь_жилого_помещения_м2 = ?", ([valueForHousingStock]))
            newValueForTableHousingStock = database.cur.fetchall()
            print(newValueForTableHousingStock)
            if newValueForTableHousingStock == []:
                print("Этот не Площадь жилого помещения")
            else:
                self.tableHousingStock.clearContents()
                self.tableHousingStock.setRowCount(len(newValueForTableHousingStock))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][0])))
                if i == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][1])))
                if i == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][2])))
                if i == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][3])))
                if i == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][4])))
                if i == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableHousingStock = []
            database.cur.execute("SELECT * FROM 'housing stock' WHERE Регистрация = ?", ([valueForHousingStock]))
            newValueForTableHousingStock = database.cur.fetchall()
            print(newValueForTableHousingStock)
            if newValueForTableHousingStock == []:
                print("Этот не Регистрация")
            else:
                self.tableHousingStock.clearContents()
                self.tableHousingStock.setRowCount(len(newValueForTableHousingStock))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][0])))
                if i == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][1])))
                if i == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][2])))
                if i == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 3, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][3])))
                if i == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableHousingStock)):
                self.tableHousingStock.setItem(i, 4, QtWidgets.QTableWidgetItem(str(newValueForTableHousingStock[i][4])))
                if i == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

    def searchForStorage(self):
        valueForStorage = self.lineForSearchingForStorage.text()

        try:
            newValueForTableStorage = []
            database.cur.execute("SELECT * FROM storage WHERE Наименование = ?", ([valueForStorage]))
            newValueForTableStorage = database.cur.fetchall()
            print(newValueForTableStorage)
            if newValueForTableStorage == []:
                print("Этот не Наименование")
            else:
                self.tableStorage.clearContents()
                self.tableStorage.setRowCount(len(newValueForTableStorage))

            for i in range(len(newValueForTableStorage)):
                self.tableStorage.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableStorage[i][0])))
                if i == 0:
                    self.tableStorage.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStorage.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStorage)):
                self.tableStorage.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableStorage[i][1])))
                if i == 0:
                    self.tableStorage.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStorage.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableStorage)):
                self.tableStorage.setItem(i, 2, QtWidgets.QTableWidgetItem(str(newValueForTableStorage[i][2])))
                if i == 0:
                    self.tableStorage.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStorage.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

    def searchForResidents(self):
        valueForResidents = self.lineForSearchingForResidents.text()

        try:
            newValueForTableResidents = []
            database.cur.execute("SELECT * FROM residents WHERE ФИО = ?", ([valueForResidents]))
            newValueForTableResidents = database.cur.fetchall()
            print(newValueForTableResidents)
            if newValueForTableResidents == []:
                print("Этот не ФИО")
            else:
                self.tableResidents.clearContents()
                self.tableResidents.setRowCount(len(newValueForTableResidents))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][5]))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableResidents = []
            database.cur.execute("SELECT * FROM residents WHERE Пол = ?", ([valueForResidents]))
            newValueForTableResidents = database.cur.fetchall()
            print(newValueForTableResidents)
            if newValueForTableResidents == []:
                print("Этот не Пол")
            else:
                self.tableResidents.clearContents()
                self.tableResidents.setRowCount(len(newValueForTableResidents))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][5]))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableResidents = []
            database.cur.execute("SELECT * FROM residents WHERE Воинская_обязанность = ?", ([valueForResidents]))
            newValueForTableResidents = database.cur.fetchall()
            print(newValueForTableResidents)
            if newValueForTableResidents == []:
                print("Этот не Воинская обязанность")
            else:
                self.tableResidents.clearContents()
                self.tableResidents.setRowCount(len(newValueForTableResidents))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][5]))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableResidents = []
            database.cur.execute("SELECT * FROM residents WHERE Паспортные_данные = ?", ([valueForResidents]))
            newValueForTableResidents = database.cur.fetchall()
            print(newValueForTableResidents)
            if newValueForTableResidents == []:
                print("Этот не Паспортные данные")
            else:
                self.tableResidents.clearContents()
                self.tableResidents.setRowCount(1)

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][5]))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableResidents = []
            database.cur.execute("SELECT * FROM residents WHERE СНИЛС = ?", ([valueForResidents]))
            newValueForTableResidents = database.cur.fetchall()
            print(newValueForTableResidents)
            print(newValueForTableResidents)
            if newValueForTableResidents == []:
                print("Этот не СНИЛС")
            else:
                self.tableResidents.clearContents()
                self.tableResidents.setRowCount(len(newValueForTableResidents))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][5]))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableResidents = []
            database.cur.execute("SELECT * FROM residents WHERE №_подъезда = ?", ([valueForResidents]))
            newValueForTableResidents = database.cur.fetchall()
            print(newValueForTableResidents)
            if newValueForTableResidents == []:
                print("Этот не №_подъезда")
            else:
                self.tableResidents.clearContents()
                self.tableResidents.setRowCount(len(newValueForTableResidents))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][5]))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableResidents = []
            database.cur.execute("SELECT * FROM residents WHERE №_квартиры = ?", ([valueForResidents]))
            newValueForTableResidents = database.cur.fetchall()
            print(newValueForTableResidents)
            if newValueForTableResidents == []:
                print("Этот не №_квартиры")
            else:
                self.tableResidents.clearContents()
                self.tableResidents.setRowCount(len(newValueForTableResidents))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][5]))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)

        try:
            newValueForTableResidents = []
            database.cur.execute("SELECT * FROM residents WHERE Задолженность = ?", ([valueForResidents]))
            newValueForTableResidents = database.cur.fetchall()
            print(newValueForTableResidents)
            if newValueForTableResidents == []:
                print("Этот не Задолженность")
            else:
                self.tableResidents.clearContents()
                self.tableResidents.setRowCount(len(newValueForTableResidents))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][5]))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(newValueForTableResidents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(newValueForTableResidents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(newValueForTableResidents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except Exception as e:
            print(e)


    def parsing(self):
        self.pars_window.show()

    def Calculating(self):
        try:
            water = int(self.lineEdit_6.text())
            water_disposal = int(self.lineEdit_7.text())
            electricity = int(self.lineEdit_5.text())

            summ = int(water * 22.46 + water_disposal * 20.20 + electricity * 4.53)
            self.SumCommonArea.setText(str(summ) + " руб.")

            database.cur.execute("SELECT * FROM 'utilities'")
            self.lenUtilities = database.cur.fetchall()
            print(len(self.lenUtilities))

            self.newUtilities = []
            self.newUtilities.append(str(len(self.lenUtilities) + 1))
            self.newUtilities.append(self.comboForCommonArea.currentText())
            self.newUtilities.append(str("Оплата МОП"))
            self.newUtilities.append(str(datetime.datetime.now().strftime('%Y.%m.%d')))
            self.newUtilities.append(str(electricity))
            self.newUtilities.append(str(water))
            self.newUtilities.append(str(water_disposal))
            self.newUtilities.append(str(summ))
            database.cur.execute("INSERT OR IGNORE INTO 'utilities' VALUES(?,?,?,?,?,?,?,?)", self.newUtilities)
            database.conn.commit()
            self.insert_data_for_utilites()

        except:
            print("ERROR")

    def treeDocs(self, item):
        print(item.text(0))
        self.listForDocs.clear()
        if item.text(0) == "Проведение работ":
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/work")
            for i in range(len(directory)):
                self.listForDocs.insertItem(i+1, str(directory[i]))
        elif item.text(0) == "Выполненные":
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/completed")
            for i in range(len(directory)):
                self.listForDocs.insertItem(i+1, str(directory[i]))
        elif item.text(0) == "Невыполненные":
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/unfulfilled")
            for i in range(len(directory)):
                self.listForDocs.insertItem(i+1, str(directory[i]))
        elif item.text(0) == "Ожидающие":
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/waiting")
            for i in range(len(directory)):
                self.listForDocs.insertItem(i+1, str(directory[i]))
        elif item.text(0) == "Коммунальные услуги":
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/service")
            for i in range(len(directory)):
                self.listForDocs.insertItem(i+1, str(directory[i]))
        elif item.text(0) == "СУиРОИ":
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/serviceCommonArea")
            for i in range(len(directory)):
                self.listForDocs.insertItem(i+1, str(directory[i]))
        elif item.text(0) == "Жильцам":
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/forResidents")
            for i in range(len(directory)):
                self.listForDocs.insertItem(i+1, str(directory[i]))

    def CheckableApplications(self, index):
        self.listWidgetForCheckable.clear()
        if index == 0:
            print("Выполненные")
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/completed")
            for i in range(len(directory)):
                self.listWidgetForCheckable.insertItem(0, str(directory[i]))
        elif index == 1:
            print("Невыполненные")
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/unfulfilled")
            for i in range(len(directory)):
                self.listWidgetForCheckable.insertItem(0, str(directory[i]))
        elif index == 2:
            print("Ожидающие")
            directory = os.listdir("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/waiting")
            for i in range(len(directory)):
                self.listWidgetForCheckable.insertItem(0, str(directory[i]))

    def showTime(self):
        #currentTime = QTime.currentTime()
        #text_time = currentTime.toString("hh:mm:ss")
        dt = datetime.datetime.now()
        self.statusbar.showMessage(str(dt.strftime('%Y.%m.%d - %H:%M:%S')), 1000)

    def addForWorkStatement(self):
        self.newWorkStatement = []
        self.newWorkStatement.append(len(self.result_work_statement) + 1)
        self.newWorkStatement.append(self.lineEditStatisticForStorage1.text())
        self.newWorkStatement.append('{}.{}.{}'.format(self.dateEdit.dateTime().toString('dd'), self.dateEdit.dateTime().toString('MM'), self.dateEdit.dateTime().toString('yyyy')))
        self.newWorkStatement.append("-")
        self.newWorkStatement.append("-")
        print(self.newWorkStatement)

        try:
            database.cur.execute("INSERT OR IGNORE INTO 'works statement' VALUES(?,?,?,?,?);", self.newWorkStatement)
            database.conn.commit()

            print(f"{datetime.datetime.now()} Запись добавлена в таблицу \"works statement\"")
            self.createWorkStatementDocument(self.lineEditStatisticForStorage1.text(), '{}.{}.{}'.format(self.dateEdit.dateTime().toString('dd'), self.dateEdit.dateTime().toString('MM'), self.dateEdit.dateTime().toString('yyyy')))

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"works statement\"")

    def openFileForStatement(self, item):
        if self.comboType_2.currentIndex() == 0:
            os.startfile("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/completed/" + item.text())
        elif self.comboType_2.currentIndex() == 1:
            os.startfile("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/unfulfilled/" + item.text())
        elif self.comboType_2.currentIndex() == 2:
            os.startfile("C:/Users/piplofen/Desktop/диплом vol.2/docs/request/waiting/" + item.text())

    def openFileForStorage(self, item):
        os.startfile("C:/Users/piplofen/Desktop/диплом vol.2/docs/work/" + item.text())

    def delete_data_for_storage(self):#ДОДЕЛАТЬ УДАЛЕНИЕ НЕСКОЛЬКИХ ШТУК
        itemName = self.comboForNameStorageOutgoing.currentText()
        print(itemName)
        print(self.lineForCountStorageOutgoing.text())

        try:
            database.cur.execute("DELETE FROM 'storage' WHERE Наименование = ?", (itemName,))
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись удалена из таблицы \"storage\"")
            self.insert_data_for_storage()

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не удалена из таблицы \"storage\"")

    def AddForStorage(self):
        self.newStorage = []
        self.newStorage.append(len(self.result_storage) + 2)
        self.newStorage.append(self.lineForNameStorageArrival.text())
        self.newStorage.append(self.lineForCountStorageArrival.text())
        print(self.newStorage)

        try:
            database.cur.execute("INSERT OR IGNORE INTO 'storage' VALUES(?,?,?)", self.newStorage)
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись добавлена в таблицу \"storage\"")
            self.insert_data_for_storage()

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"storage\"")


    def delRowForResidents(self):
        item = self.tableResidents.currentRow() + 1
        print(item)

        try:
            database.cur.execute("DELETE FROM 'residents' WHERE ID = ?", (item,))
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись удалена из таблицы \"residents\"")
            self.insert_data_for_residents()

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не удалена из таблицы \"residents\"")

    def addResidents(self):
    
        try:
            item = self.tableResidents.currentRow()
            print(item)
            data = []
            for i in range(self.tableResidents.columnCount()):
                data.append(self.tableResidents.item(item, i).text())

            database.cur.execute("INSERT OR IGNORE INTO 'residents' VALUES(?,?,?,?,?,?,?,?,?)", data)
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись добавлена в таблицу \"residents\"")

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"residents\"")

    def addRowForResidents(self):
        lastItem = self.tableResidents.rowCount()
        self.tableResidents.insertRow(lastItem)
        lastItem = 0

    def delRowForHousingStock(self):
        item = self.tableHousingStock.currentRow() + 1
        print(item)

        try:
            database.cur.execute("DELETE FROM 'housing stock' WHERE ID = ?", (item,))
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись удалена из таблицы \"housing stock\"")
            self.insert_data_for_housing_stock()

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не удалена из таблицы \"housing stock\"")

    def addRowForHousingStock(self):
        lastItem = self.tableHousingStock.rowCount()
        self.tableHousingStock.insertRow(lastItem)
        lastItem = 0

    def addHousingStock(self):
        item = self.tableHousingStock.currentRow()
        print(item)
        data = []
        for i in range(self.tableHousingStock.columnCount()):
            data.append(self.tableHousingStock.item(item, i).text())
        print(data)

        try:
            database.cur.execute("INSERT OR IGNORE INTO 'housing stock' VALUES(?,?,?,?,?)", data)
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись добавлена в таблицу \"housing stock\"")

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"housing stock\"")


    def delRowForCommonAreas(self):
        item = self.tableCommonAreas.currentRow() + 1
        print(item)

        try:
            database.cur.execute("DELETE FROM 'common areas' WHERE ID = ?", (item,))
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись удалена из таблицы \"housing stock\"")
            self.insert_data_for_common_areas()

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не удалена из таблицы \"housing stock\"")

    def addRowForCommonAreas(self):
        lastItem = self.tableCommonAreas.rowCount()
        self.tableCommonAreas.insertRow(lastItem)
        lastItem = 0

    def addCommonAreas(self):
        item = self.tableCommonAreas.currentRow()
        print(item)
        data = []
        for i in range(self.tableCommonAreas.columnCount()):
            data.append(self.tableCommonAreas.item(item, i).text())
        print(data)

        try:
            database.cur.execute("INSERT OR IGNORE INTO 'common areas' VALUES(?,?,?)", data)
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись добавлена в таблицу \"common areas\"")

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"common areas\"")

    def delRowForDebtors(self):
        item = self.tableDebtors.currentRow() + 1
        print(item)

        try:
            database.cur.execute("DELETE FROM 'debtors' WHERE ID = ?", (item,))
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись удалена из таблицы \"debtors\"")
            self.insert_data_for_debtors()

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не удалена из таблицы \"debtors\"")

    def addRowForDebtors(self):
        lastItem = self.tableDebtors.rowCount()
        self.tableDebtors.insertRow(lastItem)
        lastItem = 0

    def addDebtors(self):
        item = self.tableDebtors.currentRow()
        print(item)
        data = []

        for i in range(self.tableDebtors.columnCount()):
            data.append(self.tableDebtors.item(item, i).text())
        print(data)

        try:
            database.cur.execute("INSERT OR IGNORE INTO 'debtors' VALUES(?,?,?,?)", data)
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись добавлена в таблицу \"debtors\"")

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"debtors\"")

    def delRowForSuppliers(self):
        item = self.tableSuppliers.currentRow() + 1
        print(item)

        try:
            database.cur.execute("DELETE FROM 'suppliers' WHERE ID = ?", (item,))
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись удалена из таблицы \"suppliers\"")
            self.insert_data_for_suppliers()

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не удалена из таблицы \"suppliers\"")

    def addRowForSuppliers(self):
        lastItem = self.tableSuppliers.rowCount()
        self.tableSuppliers.insertRow(lastItem)
        lastItem = 0

    def addSuppliers(self):
        item = self.tableSuppliers.currentRow()
        print(item)
        data = []

        for i in range(self.tableSuppliers.columnCount()):
            data.append(self.tableSuppliers.item(item, i).text())
        print(data)

        try:
            database.cur.execute("INSERT OR IGNORE INTO 'suppliers' VALUES(?,?,?,?,?,?,?,?,?,?)", data)
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись добавлена в таблицу \"suppliers\"")

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"suppliers\"")

    def addStatement(self):
        newStatement = []
        newStatement.append(len(self.result_statement) + 1)
        newStatement.append(self.comboName.currentText())
        newStatement.append(self.comboType.currentText())
        newStatement.append('{}.{}.{}'.format(self.dateEdit_2.dateTime().toString('dd'), self.dateEdit_2.dateTime().toString('MM'), self.dateEdit_2.dateTime().toString('yyyy')))

        try:
            database.cur.execute("INSERT OR IGNORE INTO 'statement' VALUES(?,?,?,?)", newStatement)
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись добавлена в таблицу \"statement\"")
            self.insert_data_for_statement()
            self.createStatementDocument(self.comboName.currentText(), self.comboType.currentText(), '{}.{}.{}'.format(self.dateEdit_2.dateTime().toString('dd'), self.dateEdit_2.dateTime().toString('MM'), self.dateEdit_2.dateTime().toString('yyyy')))

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"statement\"")

    def delRowForStatement(self):
        item = self.tableStatement.currentRow() + 1
        print(item)

        try:
            database.cur.execute("DELETE FROM 'statement' WHERE ID = ?", (item,))
            database.conn.commit()
            print(f"{datetime.datetime.now()} Запись удалена из таблицы \"statement\"")
            self.insert_data_for_statement()

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не удалена из таблицы \"statement\"")

    def insert_data_for_utilites(self):
        database.cur.execute("SELECT * FROM utilities;")
        
        self.result_utilities = database.cur.fetchall()
        self.tableUtilites.setColumnWidth(0, 10)
        self.tableUtilites.setColumnWidth(1, 130)
        self.tableUtilites.setColumnWidth(2, 130)
        self.tableUtilites.setColumnWidth(3, 130)
        self.tableUtilites.setColumnWidth(4, 90)
        self.tableUtilites.setColumnWidth(5, 130)
        self.tableUtilites.setColumnWidth(6, 120)
        self.tableUtilites.verticalHeader().setVisible(False)

        try:
            self.tableUtilites.setRowCount(len(self.result_utilities))

            for i in range(len(self.result_utilities)):
                self.tableUtilites.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_utilities[i][0])))
                if i == 0:
                    self.tableUtilites.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableUtilites.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_utilities)):
                self.tableUtilites.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_utilities[i][1])))
                if i == 0:
                    self.tableUtilites.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableUtilites.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_utilities)):
                self.tableUtilites.setItem(i, 2, QtWidgets.QTableWidgetItem(str(self.result_utilities[i][3])))
                if i == 0:
                    self.tableUtilites.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableUtilites.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_utilities)):
                self.tableUtilites.setItem(i, 3, QtWidgets.QTableWidgetItem(str(self.result_utilities[i][4])))
                if i == 0:
                    self.tableUtilites.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableUtilites.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_utilities)):
                self.tableUtilites.setItem(i, 4, QtWidgets.QTableWidgetItem(str(self.result_utilities[i][5])))
                if i == 0:
                    self.tableUtilites.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableUtilites.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_utilities)):
                self.tableUtilites.setItem(i, 5, QtWidgets.QTableWidgetItem(str(self.result_utilities[i][6])))
                if i == 0:
                    self.tableUtilites.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableUtilites.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_utilities)):
                self.tableUtilites.setItem(i, 6, QtWidgets.QTableWidgetItem(str(self.result_utilities[i][7])))
                if i == 0:
                    self.tableUtilites.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableUtilites.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))



        except:
            print("123")

    def insert_data_for_residents(self):
        self.lineForSearchingForResidents.setText('')
        database.cur.execute("SELECT * FROM residents;")
        self.result_residents = database.cur.fetchall()
        self.tableResidents.setColumnWidth(0, 10)
        self.tableResidents.setColumnWidth(1, 320)
        self.tableResidents.setColumnWidth(2, 60)
        self.tableResidents.setColumnWidth(3, 225)
        self.tableResidents.setColumnWidth(4, 215)
        self.tableResidents.setColumnWidth(5, 130)
        self.tableResidents.setColumnWidth(6, 130)
        self.tableResidents.setColumnWidth(7, 130)
        self.tableResidents.setColumnWidth(8, 150)
        self.tableResidents.verticalHeader().setVisible(False)
        print(self.result_residents)

        try:
            #self.TableResidents.clearContents()
            self.tableResidents.setRowCount(len(self.result_residents))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_residents[i][0])))
                if i == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_residents[i][1])))
                if i == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 2, QtWidgets.QTableWidgetItem(self.result_residents[i][2]))
                if i == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 3, QtWidgets.QTableWidgetItem(self.result_residents[i][3]))
                if i == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 4, QtWidgets.QTableWidgetItem(self.result_residents[i][4]))
                if i == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 5, QtWidgets.QTableWidgetItem(str(self.result_residents[i][5])))
                if i == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 6, QtWidgets.QTableWidgetItem(str(self.result_residents[i][6])))
                if i == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 7, QtWidgets.QTableWidgetItem(str(self.result_residents[i][7])))
                if i == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_residents)):
                self.tableResidents.setItem(i, 8, QtWidgets.QTableWidgetItem(self.result_residents[i][8]))
                if i == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableResidents.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

        except AttributeError as e:
            print(f"{datetime.datetime.now()} Ошибка загрузки таблицы \"residents\"" + e)

    def insert_data_for_storage(self):
        self.lineForSearchingForStorage.setText('')
        database.cur.execute("SELECT * FROM 'storage'")
        self.result_storage = database.cur.fetchall()
        self.tableStorage.setColumnWidth(0, 50)
        self.tableStorage.setColumnWidth(1, 250)
        self.tableStorage.setColumnWidth(2, 150)
        self.tableStorage.verticalHeader().setVisible(False)

        try:
            self.tableStorage.setRowCount(len(self.result_storage))

            for i in range(len(self.result_storage)):
                self.tableStorage.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_storage[i][0])))
                if i == 0:
                    self.tableStorage.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStorage.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_storage)):
                self.tableStorage.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_storage[i][1])))
                if i == 0:
                    self.tableStorage.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStorage.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_storage)):
                self.tableStorage.setItem(i, 2, QtWidgets.QTableWidgetItem(str(self.result_storage[i][2])))
                if i == 0:
                    self.tableStorage.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStorage.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"storage\"")


    def insert_data_for_work_statement(self):
        self.comboForWorkStatement.clear()
        database.cur.execute("SELECT Тип_работ FROM 'works statement';")
        self.result_type_work_statement = database.cur.fetchall()
        print(self.result_type_work_statement)
        for i in range(len(self.result_type_work_statement)):
            self.comboForWorkStatement.addItems(self.result_type_work_statement[i])

        self.lineForSearchingForWorkStatement.setText('')
        database.cur.execute("SELECT * FROM 'works statement'")
        self.result_work_statement = database.cur.fetchall()
        self.tableWorkStatement.setColumnWidth(0, 1)
        self.tableWorkStatement.setColumnWidth(1, 300)
        self.tableWorkStatement.setColumnWidth(2, 150)
        self.tableWorkStatement.setColumnWidth(3, 150)
        self.tableWorkStatement.setColumnWidth(4, 150)
        self.tableWorkStatement.verticalHeader().setVisible(False)

        try:
            self.tableWorkStatement.setRowCount(len(self.result_work_statement))

            for i in range(len(self.result_work_statement)):
                self.tableWorkStatement.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_work_statement[i][0])))
                if i == 0:
                    self.tableWorkStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_work_statement)):
                self.tableWorkStatement.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_work_statement[i][1])))
                if i == 0:
                    self.tableWorkStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_work_statement)):
                self.tableWorkStatement.setItem(i, 2, QtWidgets.QTableWidgetItem(str(self.result_work_statement[i][2])))
                if i == 0:
                    self.tableWorkStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_work_statement)):
                self.tableWorkStatement.setItem(i, 3, QtWidgets.QTableWidgetItem(str(self.result_work_statement[i][3])))
                if i == 0:
                    self.tableWorkStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_work_statement)):
                self.tableWorkStatement.setItem(i, 4, QtWidgets.QTableWidgetItem(str(self.result_work_statement[i][4])))
                if i == 0:
                    self.tableWorkStatement.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableWorkStatement.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

        except TypeError:
            print(f"{datetime.datetime.now()} Запись не добавлена в таблицу \"works statement\"")

    def insert_data_for_housing_stock(self):
        self.lineForSearchingForHousingStock.setText('')
        database.cur.execute("SELECT * FROM 'housing stock';")
        self.result_housing_stock = database.cur.fetchall()
        self.tableHousingStock.setColumnWidth(0, 10)
        self.tableHousingStock.setColumnWidth(1, 275)
        self.tableHousingStock.setColumnWidth(2, 320)
        self.tableHousingStock.setColumnWidth(3, 370)
        self.tableHousingStock.setColumnWidth(4, 280)
        self.tableHousingStock.verticalHeader().setVisible(False)

        try:
            #self.TableResidents.clearContents()
            self.tableHousingStock.setRowCount(len(self.result_housing_stock))

            for i in range(len(self.result_housing_stock)):
                self.tableHousingStock.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_housing_stock[i][0])))
                if i == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_housing_stock)):
                self.tableHousingStock.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_housing_stock[i][1])))
                if i == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_housing_stock)):
                self.tableHousingStock.setItem(i, 2, QtWidgets.QTableWidgetItem(str(self.result_housing_stock[i][2])))
                if i == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_housing_stock)):
                self.tableHousingStock.setItem(i, 3, QtWidgets.QTableWidgetItem(str(self.result_housing_stock[i][3])))
                if i == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_housing_stock)):
                self.tableHousingStock.setItem(i, 4, QtWidgets.QTableWidgetItem(str(self.result_housing_stock[i][4])))
                if i == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableHousingStock.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

        except AttributeError:
            print(f"{datetime.datetime.now()} Ошибка загрузки таблицы \"housing stock\"")

    def insert_data_for_common_areas(self):
        self.lineForSearchingForCommonAreas.setText('')
        database.cur.execute("SELECT * FROM 'common areas';")
        self.result_common_areas = database.cur.fetchall()
        self.tableCommonAreas.setColumnWidth(0, 10)
        self.tableCommonAreas.setColumnWidth(1, 250)
        self.tableCommonAreas.setColumnWidth(2, 250)
        self.tableCommonAreas.verticalHeader().setVisible(False)

        try:
            #self.TableResidents.clearContents()
            self.tableCommonAreas.setRowCount(len(self.result_common_areas))

            for i in range(len(self.result_common_areas)):
                self.tableCommonAreas.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_common_areas[i][0])))
                if i == 0:
                    self.tableCommonAreas.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_common_areas)):
                self.tableCommonAreas.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_common_areas[i][1])))
                if i == 0:
                    self.tableCommonAreas.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_common_areas)):
                self.tableCommonAreas.setItem(i, 2, QtWidgets.QTableWidgetItem(str(self.result_common_areas[i][2])))
                if i == 0:
                    self.tableCommonAreas.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableCommonAreas.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

        except AttributeError:
            print(f"{datetime.datetime.now()} Ошибка загрузки таблицы \"common areas\"")

    def insert_data_for_debtors(self):
        self.lineForSearchingForDebtors.setText('')
        database.cur.execute("SELECT * FROM 'debtors';")
        self.result_debtors = database.cur.fetchall()
        self.tableDebtors.setColumnWidth(0, 10)
        self.tableDebtors.setColumnWidth(1, 275)
        self.tableDebtors.setColumnWidth(2, 250)
        self.tableDebtors.setColumnWidth(3, 250)
        self.tableDebtors.verticalHeader().setVisible(False)

        try:
            #self.TableResidents.clearContents()
            self.tableDebtors.setRowCount(len(self.result_debtors))

            for i in range(len(self.result_debtors)):
                self.tableDebtors.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_debtors[i][0])))
                if i == 0:
                    self.tableDebtors.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_debtors)):
                self.tableDebtors.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_debtors[i][1])))
                if i == 0:
                    self.tableDebtors.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_debtors)):
                self.tableDebtors.setItem(i, 2, QtWidgets.QTableWidgetItem(str(self.result_debtors[i][2])))
                if i == 0:
                    self.tableDebtors.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_debtors)):
                self.tableDebtors.setItem(i, 3, QtWidgets.QTableWidgetItem(str(self.result_debtors[i][3])))
                if i == 0:
                    self.tableDebtors.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableDebtors.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

        except AttributeError:
            print(f"{datetime.datetime.now()} Ошибка загрузки таблицы \"debtors\"")

    def insert_data_for_suppliers(self):
        self.lineForSearchingForSuppliers.setText('')
        database.cur.execute("SELECT * FROM 'suppliers'")
        self.result_suppliers = database.cur.fetchall()
        self.tableSuppliers.setColumnWidth(0, 50)
        self.tableSuppliers.setColumnWidth(1, 300)
        self.tableSuppliers.setColumnWidth(2, 250)
        self.tableSuppliers.setColumnWidth(3, 100)
        self.tableSuppliers.setColumnWidth(4, 100)
        self.tableSuppliers.setColumnWidth(5, 200)
        self.tableSuppliers.setColumnWidth(6, 200)
        self.tableSuppliers.setColumnWidth(7, 100)
        self.tableSuppliers.setColumnWidth(8, 100)
        self.tableSuppliers.setColumnWidth(9, 200)
        self.tableSuppliers.verticalHeader().setVisible(False)

        try:
            self.tableSuppliers.setRowCount(len(self.result_suppliers))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][0])))
                if i == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][1])))
                if i == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 2, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][2])))
                if i == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 3, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][3])))
                if i == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 4, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][4])))
                if i == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 4).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 5, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][5])))
                if i == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 5).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 6, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][6])))
                if i == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 6).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 7, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][7])))
                if i == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 7).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 8, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][8])))
                if i == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 8).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_suppliers)):
                self.tableSuppliers.setItem(i, 9, QtWidgets.QTableWidgetItem(str(self.result_suppliers[i][9])))
                if i == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableSuppliers.item(i, 9).setBackground(QBrush(QColor(220, 220, 220)))

        except AttributeError:
            print(f"{datetime.datetime.now()} Ошибка загрузки таблицы \"suppliers\"")

    def insert_data_for_statement(self):
        self.lineForSearchingForStatement.setText('')
        database.cur.execute("SELECT * FROM 'statement';")
        self.result_statement = database.cur.fetchall()
        self.tableStatement.setColumnWidth(0, 50)
        self.tableStatement.setColumnWidth(1, 275)
        self.tableStatement.setColumnWidth(2, 250)
        self.tableStatement.setColumnWidth(3, 250)
        self.tableStatement.verticalHeader().setVisible(False)

        try:
            #self.TableResidents.clearContents()
            self.tableStatement.setRowCount(len(self.result_statement))

            for i in range(len(self.result_statement)):
                self.tableStatement.setItem(i, 0, QtWidgets.QTableWidgetItem(str(self.result_statement[i][0])))
                if i == 0:
                    self.tableStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 0).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_statement)):
                self.tableStatement.setItem(i, 1, QtWidgets.QTableWidgetItem(str(self.result_statement[i][1])))
                if i == 0:
                    self.tableStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 1).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_statement)):
                self.tableStatement.setItem(i, 2, QtWidgets.QTableWidgetItem(str(self.result_statement[i][2])))
                if i == 0:
                    self.tableStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 2).setBackground(QBrush(QColor(220, 220, 220)))

            for i in range(len(self.result_statement)):
                self.tableStatement.setItem(i, 3, QtWidgets.QTableWidgetItem(str(self.result_statement[i][3])))
                if i == 0:
                    self.tableStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))
                elif i % 2 == 0:
                    self.tableStatement.item(i, 3).setBackground(QBrush(QColor(220, 220, 220)))

        except AttributeError:
            print(f"{datetime.datetime.now()} Ошибка загрузки таблицы \"statement\"")

def main():
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    window.show()
    app.exec_()

if __name__ == '__main__':#Если мы запускаем файл напрямую, а не импортируем
    main()#то запускаем функцию main()